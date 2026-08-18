#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""简历内容分类脚本：读取简历正文，按内容关键词归入职位类目

用法:
    classify_content.py <目录> --out report.json            预演
    classify_content.py <目录> --apply --target <根目录>    移动

评分规则：每个类目一组关键词，全文命中计数，取最高分类；
最高分 >= MIN_SCORE 才移动，否则留在原目录（未分类）。

Created & maintained by Glen Wei (韦其像)
Email: glen.keeming@gmail.com
Part of headhunter-skills: https://github.com/Glen-Wei/headhunter-skills"""
import os, re, sys, json, shutil

AUTHOR_EPILOG = (
    "Author: Glen Wei (韦其像) | Email: glen.keeming@gmail.com | "
    "Part of headhunter-skills: https://github.com/Glen-Wei/headhunter-skills"
)

MIN_SCORE = 2

CONTENT_RULES = {
    "强化学习": ["强化学习", "reinforcement learning", "RLHF", "PPO", "DDPG", "SAC", "reward", "运控", "运动控制", "Sim2Real", "sim-to-real", "足式机器人", "双足"],
    "机械设计": ["机械设计", "结构设计", "SolidWorks", "CAD", "Creo", "Catia", "有限元", "FEA", "Ansys", "Abaqus", "公差", "尺寸链", "齿轮", "减速器", "机械结构"],
    "SLAM": ["SLAM", "slam", "定位建图", "激光雷达", "Lidar", "ORB-SLAM", "Cartographer", "IMU", "里程计", "建图"],
    "VLA": ["VLA", "视觉语言模型", "vision-language", "VLM", "多模态大模型", "CLIP", "具身智能体", "embodied", "动作生成", "语言指令"],
    "LLM": ["大模型", "LLM", "GPT", "LLaMA", "Qwen", "NLP", "自然语言处理", "prompt", "微调", "fine-tuning", "transformer", "预训练", "ChatGLM"],
    "仿真": ["仿真", "Simulation", "Isaac", "Gazebo", "MuJoCo", "Mujoco", "PyBullet", "物理引擎", "数字孪生", "simulation", "Unity", "虚幻引擎"],
    "多模态": ["多模态", "multimodal", "多模态大模型", "跨模态", "图文", "音频", "video understanding", "视频理解"],
    "数据": ["数据标注", "数据采集", "数据清洗", "数据集", "dataset", "data pipeline", "数据管线", "爬虫", "数据工程师", "标注"],
    "商业管理": ["销售", "商务", "BD", "市场", "品牌", "客户", "大客户", "渠道", "营收", "商业化", "业务拓展", "sales"],
    "公关PR": ["公关", "PR", "媒介", "舆情", "品牌公关", "媒体", "新闻", "发布会", "危机公关", "党央媒"],
    "项目经理": ["产品经理", "项目经理", "项目管理", "PM", "TPM", "产品规划", "需求分析", "roadmap", "交付", "项目负责人", "PMP"],
    "嵌入式": ["嵌入式", "嵌入式软件", "RTOS", "单片机", "STM32", "ARM", "驱动", "Linux内核", "固件", "firmware", "ESP32"],
    "硬件": ["硬件", "电路设计", "PCB", "原理图", "layout", "电源设计", "信号完整性", "FPGA", "硬件工程师"],
    "测试": ["测试开发", "自动化测试", "测试工程师", "QA", "软件测试", "性能测试", "测试用例", "pytest", "Selenium", "Jira"],
    "供应链": ["供应链", "采购", "供应商", "物流", "库存", "产能", "supply chain", "Sourcing", "量产管理"],
    "技术美术": ["技术美术", "TA", "Shader", "材质", "渲染管线", "Houdini", "Substance", "数字雕刻", "3D模型", "贴图", "Unity", "Unreal", "UE5"],
    "算法": ["深度学习", "机器学习", "PyTorch", "TensorFlow", "目标检测", "图像分割", "神经网络", "CNN", "训练", "模型", "算法工程师", "transformer"],
    "高管": ["CEO", "CTO", "CFO", "COO", "创始人", "联合创始人", "合伙人", "VP", "副总裁", "总裁", "总经理", "事业部负责人"],
    "解决方案": ["解决方案", "售前", "解决方案架构", "方案设计", "客户方案", "技术方案", "presales"],
    "运营": ["运营", "用户运营", "内容运营", "增长", "活动运营", "社群", "运营负责人"],
    "3D重建与算法": ["三维重建", "3D重建", "NeRF", "3DGS", "高斯泼溅", "点云", "mesh", "网格重建", "photogrammetry"],
    "灵巧手": ["灵巧手", "dexterous", "机械手", "五指", "末端执行器", "抓取", "grasp"],
    "感知与运动": ["目标检测", "物体识别", "位姿估计", "视觉SLAM", "光流", "深度估计", "人体姿态", "运动规划", "轨迹规划", "感知算法", "detection", "segmentation"],
    "控制规划": ["控制算法", "控制理论", "PID", "MPC", "LQR", "鲁棒控制", "轨迹跟踪", "动力学", "控制器设计", "力控"],
    "机器人部署": ["机器人部署", "部署落地", "现场部署", "实施", "运维部署", "集成", "调试", "落地交付"],
    "信息安全": ["信息安全", "网络安全", "安全工程师", "渗透测试", "等保", "加密", "SOC", "威胁", "漏洞", "security"],
    "世界模型": ["世界模型", "world model", "World Model", "world-model", "可计算空间智能"],
    "基础设施SRE": ["Infra", "基础设施", "SRE", "Kubernetes", "K8s", "Docker", "云平台", "CI/CD", "GPU集群", "分布式训练", "DevOps", "监控"],
    "Agent": ["Agent", "智能体", "多智能体", "工具调用", "function calling", "RAG", "自主决策", "agentic"],
    "具身研究员": ["博士后", "助理教授", "副教授", "教授", "研究员", "博士", "phd", "research scientist", "学术"],
    "自动驾驶": ["自动驾驶", "ADAS", "智能驾驶", "高精地图", "路径规划", "泊车", "NOA", "autonomous driving"],
}

EXTRACTORS = {}
import subprocess

def extract_pdf(path):
    try:
        from pypdf import PdfReader
        r = PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in r.pages)
    except Exception as e:
        return ""

def extract_docx(path):
    try:
        from docx import Document
        d = Document(path)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.append(" ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext in (".docx", ".doc"):
        return extract_docx(path)
    return ""

def classify_content(text):
    low = text.lower()
    scores = {}
    for cat, kws in CONTENT_RULES.items():
        n = 0
        for kw in kws:
            if kw.lower() in low:
                n += 1
        if n:
            scores[cat] = n
    if not scores:
        return None, 0
    best = max(scores, key=scores.get)
    return best, scores[best]

def main():
    args, opts = [], {}
    i = 1
    while i < len(sys.argv):
        a = sys.argv[i]
        if a.startswith("--"):
            if i + 1 < len(sys.argv) and not sys.argv[i + 1].startswith("--"):
                opts[a[2:]] = sys.argv[i + 1]; i += 2
            else:
                opts[a[2:]] = True; i += 1
        else:
            args.append(a); i += 1
    if not args:
        print(__doc__); return
    src_dir = args[0]
    target = opts.get("target")
    out = opts.get("out")
    apply_move = opts.get("apply")
    report = {"可分类": [], "低置信(留未分类)": [], "无文本层": []}
    files = sorted(os.listdir(src_dir))
    for f in files:
        if f.startswith("~$") or f == ".DS_Store":
            continue
        path = os.path.join(src_dir, f)
        if not os.path.isfile(path):
            continue
        text = extract_text(path)
        if not text.strip():
            report["无文本层"].append({"file": f})
            continue
        cat, score = classify_content(text)
        if cat and score >= MIN_SCORE:
            report["可分类"].append({"file": f, "类目": cat, "得分": score,
                                     "线索": re.sub(r"\s+", " ", text)[:120]})
        else:
            report["低置信(留未分类)"].append({"file": f, "类目": cat, "得分": score,
                                               "线索": re.sub(r"\s+", " ", text)[:120]})
    if out:
        with open(out, "w", encoding="utf-8") as fh:
            json.dump(report, fh, ensure_ascii=False, indent=1)
    for section in ("可分类", "低置信(留未分类)", "无文本层"):
        items = report[section]
        print(f"══ {section}：{len(items)} 份 ══")
        for it in items[:8]:
            extra = f" → {it['类目']}({it['得分']})" if '类目' in it and it['类目'] else ""
            print(f"  {it['file']}{extra}｜{it.get('线索', '')[:60]}")
        if len(items) > 8:
            print(f"  ... 共 {len(items)} 份")
    if apply_move:
        assert target, "--target 必填"
        n = 0
        for it in report["可分类"]:
            cat_dir = os.path.join(target, it["类目"])
            os.makedirs(cat_dir, exist_ok=True)
            src = os.path.join(src_dir, it["file"])
            dst = os.path.join(cat_dir, it["file"])
            if os.path.exists(dst):
                base, ext = os.path.splitext(it["file"])
                dst = os.path.join(cat_dir, f"{base}_2{ext}")
            shutil.move(src, dst)
            n += 1
        print(f"✔ 已移动 {n} 个文件到 {target}")

if __name__ == "__main__":
    print(AUTHOR_EPILOG)
    main()
