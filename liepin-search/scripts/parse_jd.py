"""
JD解析器 (parse_jd.py)
======================
从 JD 文档(docx/pdf/txt)中自动提取硬性要求与评分配置，输出供 liepin_deep_review.py 消费的 JD JSON。

用法:
  python parse_jd.py <jd文件> [--output <json路径>]

提取规则(按中文JD常见表述)：
  - min_edu    : 学历下限 (大专/本科/硕士/博士)，命中 "本科及以上/硕士以上/博士" 等
  - tongzhao   : 是否要求统招 ("全日制/统招" 出现即 True)
  - min_years  : 工作年限下限 ("8年以上...经验"/"3年以上经验" 等)
  - core_skills_any : 必备方向词 (任中其一才通过硬筛)，如 "电主轴|末端执行器|电动工具|灵巧手"
  - landing_min: 落地/量产证据词最低命中数 ("至少2个...落地项目" → 2)
  - skills     : 加分技能词 (评分用)
  - industry   : 行业词 (评分用)
  - cities     : 期望城市加分列表
"""
import sys, os, re, json, argparse

# 学历等级
EDU_RANK = {'大专': 1, '本科': 2, '硕士': 3, '博士': 4}

def extract_text(path):
    """读取 docx / pdf / txt 文本"""
    lower = path.lower()
    if lower.endswith('.docx'):
        import docx
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 表格内容也读
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return '\n'.join(parts)
    if lower.endswith('.pdf'):
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return '\n'.join((pg.extract_text() or '') for pg in pdf.pages)
    # txt / md / 其他
    with open(path, encoding='utf-8', errors='ignore') as f:
        return f.read()

def parse(text):
    jd = {
        'title': '',
        'hard': {},
        'core': {},
        'skills': {},
        'industry': {},
        'landing': ['量产', '落地', '工程化', '导入', '交付', '验证', '样机', '上市', '0-1', '从0到1'],
        'cities': ['上海', '苏州', '杭州', '南京', '无锡', '常州', '昆山', '宁波'],
    }
    # 标题：第一行非空
    for line in text.split('\n'):
        line = line.strip()
        if line:
            jd['title'] = line[:60]
            break

    # ---- 学历 ----
    edu_m = re.search(r'(全日制|统招)?\s*(博士(?:后)?|硕士|本科|大专)(?:及?以上|及以上学历|及以上)?', text)
    if edu_m:
        edu = edu_m.group(2).replace('博士后', '博士')
        if '本科' in edu: edu = '本科'
        if '硕士' in edu: edu = '硕士'
        if '博士' in edu: edu = '博士'
        if '大专' in edu: edu = '大专'
        jd['hard']['min_edu'] = edu
        # "及以上" 或 "硕士/博士" 但无 "及以上" 时也视为下限
        if edu_m.group(1) in ('全日制', '统招') or '全日制' in text[:2000] or '统招' in text[:2000]:
            jd['hard']['tongzhao'] = True

    # ---- 工作年限 ----
    y_m = re.search(r'(\d+)\s*年(?:以上|及以上)?(?:的)?(?:相关)?(?:机械设计|机械|工作|从业|行业)?经验', text)
    if not y_m:
        y_m = re.search(r'(\d+)\s*年以上(?:.*?)经验', text)
    if y_m:
        jd['hard']['min_years'] = int(y_m.group(1))
    else:
        y2 = re.search(r'经验(?:要求)?[：:]?\s*(\d+)\s*年', text)
        if y2:
            jd['hard']['min_years'] = int(y2.group(1))
    # 兼容 jd_score 读取的 years_min
    jd['years_min'] = jd['hard'].get('min_years', 0)

    # ---- 必备方向词（从"岗位职责/岗位要求/任职要求/任职资格/加分项"段中挑核心名词）----
    # 常见方向名词表（可按需扩充）
    CORE_WORDS = [
        '电主轴', '主轴', '末端执行器', '末端执行机构', '灵巧手', '电动工具', '打磨头',
        '执行器', '夹具', '机械臂', '机器人', '减速器', '伺服电机', '丝杠', '导轨',
        '力控', '恒力', '柔顺', '力传感器', '结构设计', '传动', 'GD&T', '公差',
        'Ansys', 'Abaqus', '有限元', '仿真',
    ]
    core_hits = []
    for w in CORE_WORDS:
        if w in text:
            core_hits.append(w)
            # 评分权重：JD里提到的词按重要度给分
            jd['core'][w] = 12 if w in ('电主轴', '末端执行器', '灵巧手', '电动工具') else 8
    # 必备技能 = 在"任职要求/岗位要求"段出现的核心词（取前6个），硬筛要求至少命中其一
    req_seg = ''
    for marker in ['任职要求', '任职资格', '岗位要求', '招聘要求', '基本要求']:
        i = text.find(marker)
        if i >= 0:
            req_seg = text[i:i+2000]
            break
    req_core = [w for w in CORE_WORDS if w in req_seg]
    if req_core:
        jd['hard']['core_skills_any'] = req_core[:6]

    # ---- 落地项目数（"至少N个...落地/量产/产品"）----
    lp = re.search(r'至少\s*(\d+)\s*个?(?:完整)?(?:的)?(?:产品|项目)?(?:.*?)?(?:落地|量产|交付|从概念到量产)', text)
    if not lp:
        lp = re.search(r'(\d+)\s*个?(?:完整)?(?:的)?(?:产品|项目)?(?:.*?)?(?:落地|量产)', text)
    if lp:
        jd['hard']['landing_min'] = int(lp.group(1))

    # ---- 加分技能（技能表）----
    SKILL_WORDS = {
        'ansys': 8, 'abaqus': 8, '有限元': 5, '仿真': 3, 'GD&T': 8, '公差': 5,
        '尺寸链': 6, '3DCS': 6, 'solidworks': 3, 'catia': 3, 'creo': 3, 'pro/e': 3,
        'ug': 3, 'matlab': 2, 'adams': 3, 'recurdyn': 3,
    }
    for w, wt in SKILL_WORDS.items():
        if w.lower() in text.lower():
            jd['skills'][w] = wt

    # ---- 行业词 ----
    IND_WORDS = ['机器人', '人形机器人', '具身智能', '工业机器人', '自动化', '机床', '磨削', '医疗', '非标', '智能硬件', '汽车']
    for w in IND_WORDS:
        if w in text:
            jd['industry'][w] = 4 if w in ('机器人', '人形机器人', '具身智能', '工业机器人') else 2

    # ---- 工作地点（第N行"工作地点"）----
    loc = re.search(r'工作地点[：:]\s*([^\n]+)', text)
    if loc:
        jd.setdefault('location', loc.group(1).strip()[:20])

    return jd

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('jd_file', help='JD文件路径 (docx/pdf/txt)')
    ap.add_argument('--output', '-o', default=None)
    args = ap.parse_args()

    text = extract_text(args.jd_file)
    jd = parse(text)
    out = args.output or (os.path.splitext(args.jd_file)[0] + '_jd.json')
    with open(out, 'w', encoding='utf-8') as f:
        json.dump(jd, f, ensure_ascii=False, indent=2)
    print(f"✅ JD 解析完成 -> {out}")
    print(json.dumps(jd, ensure_ascii=False, indent=2))

if __name__ == '__main__':
    main()
